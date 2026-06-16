import os
import subprocess
import shutil
from pathlib import Path
from ..config import settings


class ConverterService:
    @staticmethod
    def convert_office_to_pdf(input_path: str, output_dir: str) -> str:
        cmd = [
            settings.LIBREOFFICE_PATH,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            input_path,
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

        input_name = Path(input_path).stem
        output_path = os.path.join(output_dir, f"{input_name}.pdf")
        if not os.path.exists(output_path):
            raise RuntimeError(f"Output file not found: {output_path}")
        return output_path

    @staticmethod
    def convert_pdf_to_word(input_path: str, output_dir: str) -> str:
        from pdf2docx import Converter

        input_name = Path(input_path).stem
        output_path = os.path.join(output_dir, f"{input_name}.docx")

        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()
        return output_path

    @staticmethod
    def convert_pdf_to_excel(input_path: str, output_dir: str) -> str:
        import fitz
        from openpyxl import Workbook

        input_name = Path(input_path).stem
        output_path = os.path.join(output_dir, f"{input_name}.xlsx")

        doc = fitz.open(input_path)
        wb = Workbook()
        ws = wb.active
        ws.title = "PDF Content"

        row_num = 1
        for page in doc:
            tables = page.find_tables()
            if tables and tables.tables:
                for table in tables.tables:
                    for row in table.extract():
                        for col_idx, cell in enumerate(row, 1):
                            ws.cell(row=row_num, column=col_idx, value=cell)
                        row_num += 1
                    row_num += 1
            else:
                text = page.get_text()
                for line in text.strip().split("\n"):
                    if line.strip():
                        ws.cell(row=row_num, column=1, value=line.strip())
                        row_num += 1

        doc.close()
        wb.save(output_path)
        return output_path

    @staticmethod
    def convert_pdf_to_image(input_path: str, output_dir: str, target_format: str = "png") -> str:
        import fitz

        input_name = Path(input_path).stem
        doc = fitz.open(input_path)

        if len(doc) == 1:
            page = doc[0]
            pix = page.get_pixmap(dpi=200)
            output_path = os.path.join(output_dir, f"{input_name}.{target_format}")
            if target_format == "jpg":
                pix.save(output_path)
            else:
                pix.save(output_path)
        else:
            output_dir = os.path.join(output_dir, f"{input_name}_images")
            os.makedirs(output_dir, exist_ok=True)
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=200)
                output_path = os.path.join(output_dir, f"page_{i + 1}.{target_format}")
                pix.save(output_path)

        doc.close()

        if len(doc) > 1:
            zip_path = os.path.join(settings.OUTPUT_DIR, f"{input_name}_images.zip")
            shutil.make_archive(zip_path.replace(".zip", ""), "zip", output_dir)
            return zip_path

        return output_path

    @staticmethod
    def convert_pdf_to_markdown(input_path: str, output_dir: str) -> str:
        import fitz

        input_name = Path(input_path).stem
        output_path = os.path.join(output_dir, f"{input_name}.md")

        doc = fitz.open(input_path)
        md_content = []

        for i, page in enumerate(doc):
            md_content.append(f"## Page {i + 1}\n")

            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block["type"] == 0:
                    for line in block["lines"]:
                        text = ""
                        for span in line["spans"]:
                            text += span["text"]
                        if text.strip():
                            if any(span["flags"] & 2**4 for span in line["spans"]):
                                md_content.append(f"### {text.strip()}\n")
                            elif any(span["flags"] & 2**5 for span in line["spans"]):
                                md_content.append(f"## {text.strip()}\n")
                            else:
                                md_content.append(f"{text.strip()}\n")

            images = page.get_images()
            if images:
                md_content.append(f"\n*Page {i + 1} contains {len(images)} image(s)*\n")

        doc.close()

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_content))

        return output_path

    @staticmethod
    def convert_image_to_image(input_path: str, output_dir: str, target_format: str) -> str:
        from PIL import Image

        input_name = Path(input_path).stem
        output_path = os.path.join(output_dir, f"{input_name}.{target_format}")

        img = Image.open(input_path)
        if target_format in ("jpg", "jpeg") and img.mode == "RGBA":
            img = img.convert("RGB")
        img.save(output_path)
        return output_path

    @staticmethod
    def convert_image_to_pdf(input_path: str, output_dir: str) -> str:
        from PIL import Image

        input_name = Path(input_path).stem
        output_path = os.path.join(output_dir, f"{input_name}.pdf")

        img = Image.open(input_path)
        if img.mode == "RGBA":
            img = img.convert("RGB")
        img.save(output_path, "PDF", resolution=200.0)
        return output_path

    @staticmethod
    def convert_markdown_to_pdf(input_path: str, output_dir: str) -> str:
        import fitz

        input_name = Path(input_path).stem
        output_path = os.path.join(output_dir, f"{input_name}.pdf")

        with open(input_path, "r", encoding="utf-8") as f:
            md_content = f.read()

        doc = fitz.open()

        page_width = 595
        page_height = 842
        margin_left = 72
        margin_top = 72
        margin_right = 72
        margin_bottom = 72

        current_y = margin_top
        line_height = 16

        page = doc.new_page(width=page_width, height=page_height)
        text_point = fitz.Point(margin_left, current_y)

        font_sizes = {"h1": 24, "h2": 20, "h3": 16, "h4": 14, "body": 11}

        lines = md_content.split("\n")
        for line in lines:
            stripped = line.strip()

            if not stripped:
                current_y += line_height
                if current_y > page_height - margin_bottom:
                    page = doc.new_page(width=page_width, height=page_height)
                    current_y = margin_top
                continue

            font_size = font_sizes["body"]
            is_bold = False
            clean_text = stripped

            if stripped.startswith("#### "):
                font_size = font_sizes["h4"]
                is_bold = True
                clean_text = stripped[5:]
            elif stripped.startswith("### "):
                font_size = font_sizes["h3"]
                is_bold = True
                clean_text = stripped[4:]
            elif stripped.startswith("## "):
                font_size = font_sizes["h2"]
                is_bold = True
                clean_text = stripped[3:]
            elif stripped.startswith("# "):
                font_size = font_sizes["h1"]
                is_bold = True
                clean_text = stripped[2:]
            elif stripped.startswith("**") and stripped.endswith("**"):
                is_bold = True
                clean_text = stripped[2:-2]
            elif stripped.startswith("- ") or stripped.startswith("* "):
                clean_text = "  •  " + stripped[2:]
            elif stripped.startswith("> "):
                clean_text = "  │  " + stripped[2:]
                font_size = 10

            if is_bold:
                font_size += 2

            text_point = fitz.Point(margin_left, current_y)

            try:
                page.insert_text(
                    text_point,
                    clean_text,
                    fontsize=font_size,
                    fontname="helv",
                )
            except Exception:
                page.insert_text(
                    text_point,
                    clean_text,
                    fontsize=font_size,
                )

            current_y += font_size + 6

            if current_y > page_height - margin_bottom:
                page = doc.new_page(width=page_width, height=page_height)
                current_y = margin_top

        doc.save(output_path)
        doc.close()
        return output_path

    @staticmethod
    def convert_word_to_word(input_path: str, output_dir: str, target_format: str) -> str:
        input_name = Path(input_path).stem
        input_ext = Path(input_path).suffix.lstrip(".").lower()
        output_ext = target_format

        if target_format == "txt" and input_ext in ("docx", "doc"):
            from docx import Document
            doc = Document(input_path)
            output_path = os.path.join(output_dir, f"{input_name}.txt")
            with open(output_path, "w", encoding="utf-8") as f:
                for para in doc.paragraphs:
                    f.write(para.text + "\n")
            return output_path
        else:
            cmd = [
                settings.LIBREOFFICE_PATH,
                "--headless",
                "--convert-to", output_ext,
                "--outdir", output_dir,
                input_path,
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
            output_path = os.path.join(output_dir, f"{input_name}.{output_ext}")
            if not os.path.exists(output_path):
                raise RuntimeError(f"Output file not found: {output_path}")
            return output_path

    @staticmethod
    def convert_excel_to_excel(input_path: str, output_dir: str, target_format: str) -> str:
        cmd = [
            settings.LIBREOFFICE_PATH,
            "--headless",
            "--convert-to", target_format,
            "--outdir", output_dir,
            input_path,
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

        input_name = Path(input_path).stem
        output_path = os.path.join(output_dir, f"{input_name}.{target_format}")
        if not os.path.exists(output_path):
            raise RuntimeError(f"Output file not found: {output_path}")
        return output_path

    @staticmethod
    def convert_media(input_path: str, output_dir: str, target_format: str) -> str:
        input_name = Path(input_path).stem
        output_path = os.path.join(output_dir, f"{input_name}.{target_format}")

        cmd = [
            settings.FFMPEG_PATH,
            "-y",
            "-i", input_path,
        ]

        if target_format == "gif":
            cmd.extend([
                "-vf", "fps=10,scale=480:-1:flags=lanczos",
                "-loop", "0",
            ])
        elif target_format in ("mp3", "wav", "aac", "flac", "ogg"):
            cmd.extend(["-vn", "-acodec", "libmp3lame", "-q:a", "2"])
        elif target_format in ("mp4", "mkv", "avi", "mov", "flv", "webm"):
            cmd.extend(["-c:v", "libx264", "-c:a", "aac", "-strict", "experimental"])

        cmd.append(output_path)

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
        if result.returncode != 0:
            raise RuntimeError(f"FFmpeg conversion failed: {result.stderr}")

        if not os.path.exists(output_path):
            raise RuntimeError(f"Output file not found: {output_path}")
        return output_path

    @staticmethod
    def get_media_info(file_path: str) -> dict:
        cmd = [
            settings.FFMPEG_PATH,
            "-i", file_path,
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        info_str = result.stderr

        duration = 0.0
        for line in info_str.split("\n"):
            if "Duration:" in line:
                dur_str = line.split("Duration:")[1].split(",")[0].strip()
                parts = dur_str.split(":")
                if len(parts) == 3:
                    duration = float(parts[0]) * 3600 + float(parts[1]) * 60 + float(parts[2])
                break

        return {"duration": duration, "raw": info_str}
